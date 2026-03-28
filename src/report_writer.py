"""Meridant Insight — Word document export.

Generates a client-ready .docx assessment report using python-docx.
Returns bytes suitable for st.download_button().
"""

from __future__ import annotations

import io
from datetime import date
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand colours ─────────────────────────────────────────────────────────────
_NAVY      = RGBColor(0x0F, 0x27, 0x44)
_ACCENT    = RGBColor(0x25, 0x63, 0xEB)
_GRAY_MID  = RGBColor(0x37, 0x41, 0x51)
_GRAY_LIGHT = RGBColor(0x6B, 0x72, 0x80)
_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
_TABLE_HDR = RGBColor(0x0F, 0x27, 0x44)
_GREEN_BG  = RGBColor(0xCC, 0xFB, 0xF1)
_AMBER_BG  = RGBColor(0xFE, 0xF3, 0xC7)
_RED_TXT   = RGBColor(0xDC, 0x26, 0x26)

_RISK_COLORS = {
    "High":   RGBColor(0xDC, 0x26, 0x26),
    "Medium": RGBColor(0xD9, 0x77, 0x06),
    "Low":    RGBColor(0x16, 0xA3, 0x4A),
}


def _set_cell_bg(cell, rgb_hex: str) -> None:
    """Set table cell background colour (OOXML shading)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def _hdr_para(doc: Document, text: str, level: int = 1) -> None:
    """Add a styled heading paragraph."""
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = _NAVY
    run.font.bold = True


def _table_header_row(table, *headers) -> None:
    """Style the first row of a table as a navy header."""
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = hdr
        _set_cell_bg(cell, "0F2744")
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = _WHITE
        run.font.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table(doc: Document, rows: int, cols: int, col_widths: list[float] | None = None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def generate_word_report(
    client_name: str,
    engagement_name: str,
    use_case_name: str,
    consultant_name: str,
    findings_narrative: str,
    dom_scores: list[dict],
    cap_findings: list[dict],
    recommendations: list[dict],
    framework_name: str = "Meridant Matrix",
) -> bytes:
    """
    Generate a Word assessment report.

    Returns bytes suitable for st.download_button(mime='application/vnd.openxmlformats-
    officedocument.wordprocessingml.document').
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    report_date = date.today().strftime("%d %B %Y")

    # ─────────────────────────────────────────────────────────────────────────
    # 1. Cover page
    # ─────────────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Meridant Insight")
    run.font.size = Pt(11)
    run.font.color.rgb = _ACCENT
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(client_name)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(engagement_name)
    run.font.size = Pt(14)
    run.font.color.rgb = _GRAY_MID

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Use Case: {use_case_name}")
    run.font.size = Pt(11)
    run.font.color.rgb = _GRAY_LIGHT

    if consultant_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Consultant: {consultant_name}")
        run.font.size = Pt(11)
        run.font.color.rgb = _GRAY_LIGHT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report_date)
    run.font.size = Pt(11)
    run.font.color.rgb = _GRAY_LIGHT

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 2. Executive Summary
    # ─────────────────────────────────────────────────────────────────────────
    _hdr_para(doc, "Executive Summary", level=1)
    if findings_narrative:
        for para_text in findings_narrative.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.style = "Normal"
    else:
        doc.add_paragraph("No executive summary available.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 3. Domain Findings
    # ─────────────────────────────────────────────────────────────────────────
    _hdr_para(doc, "Domain Findings", level=1)
    doc.add_paragraph(
        "The table below summarises average maturity scores, targets, and gaps per domain."
    )

    if dom_scores:
        table = _add_table(doc, rows=len(dom_scores) + 1, cols=5,
                           col_widths=[2.5, 1.0, 1.0, 0.8, 1.0])
        _table_header_row(table, "Domain", "Avg Score", "Target", "Gap", "Risk")
        for i, d in enumerate(sorted(dom_scores, key=lambda x: -(x.get("gap") or (x.get("target", 3) - (x.get("avg_score") or 0))))):
            row = table.rows[i + 1]
            avg    = d.get("avg_score")
            target = d.get("target") or 3
            gap    = round((target - avg) if avg is not None else 0, 1)
            risk   = "High" if gap >= 2 else ("Medium" if gap >= 1 else "Low")
            row.cells[0].text = d.get("domain", "")
            row.cells[1].text = f"{avg:.1f}" if avg is not None else "—"
            row.cells[2].text = str(target)
            row.cells[3].text = f"{gap:.1f}"
            row.cells[4].text = risk
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
            # Colour risk cell
            risk_color = _RISK_COLORS.get(risk)
            if risk_color:
                for run in row.cells[4].paragraphs[0].runs:
                    run.font.color.rgb = risk_color
                    run.font.bold = True
    else:
        doc.add_paragraph("No domain findings available.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 4. Capability Gap Analysis
    # ─────────────────────────────────────────────────────────────────────────
    _hdr_para(doc, "Capability Gap Analysis", level=1)

    if cap_findings:
        gap_caps = [c for c in cap_findings if (c.get("gap") or 0) > 0]
        gap_caps.sort(key=lambda x: -(x.get("gap") or 0))
        if gap_caps:
            table = _add_table(doc, rows=len(gap_caps) + 1, cols=5,
                               col_widths=[2.2, 1.8, 0.8, 0.8, 0.8])
            _table_header_row(table, "Capability", "Domain", "Score", "Target", "Gap")
            for i, c in enumerate(gap_caps):
                row = table.rows[i + 1]
                avg    = c.get("avg_score")
                target = c.get("target_maturity") or 3
                gap    = c.get("gap") or 0
                row.cells[0].text = c.get("capability_name", "")
                row.cells[1].text = c.get("domain", "")
                row.cells[2].text = f"{avg:.1f}" if avg is not None else "—"
                row.cells[3].text = str(target)
                row.cells[4].text = f"{gap:.1f}"
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)
                if gap >= 1.0:
                    for run in row.cells[4].paragraphs[0].runs:
                        run.font.color.rgb = _RED_TXT
                        run.font.bold = True
        else:
            doc.add_paragraph("No capability gaps identified.")
    else:
        doc.add_paragraph("No capability findings available.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 5. Recommendations
    # ─────────────────────────────────────────────────────────────────────────
    _hdr_para(doc, "Recommendations", level=1)

    if recommendations:
        tier_labels = {
            "P1": "P1 — Foundation (Address First)",
            "P2": "P2 — Acceleration (Strengthen Next)",
            "P3": "P3 — Optimisation (Long-term Enhancement)",
        }
        for tier in ("P1", "P2", "P3"):
            tier_recs = [r for r in recommendations if r.get("priority_tier") == tier]
            if not tier_recs:
                continue
            _hdr_para(doc, tier_labels[tier], level=2)
            for rec in tier_recs:
                _hdr_para(doc, rec.get("capability_name", ""), level=3)
                p = doc.add_paragraph()
                p.add_run("Domain: ").font.bold = True
                p.add_run(f"{rec.get('domain', '')}  |  ")
                p.add_run("Gap: ").font.bold = True
                p.add_run(f"{rec.get('gap', 0):.1f}  |  ")
                p.add_run("Effort: ").font.bold = True
                p.add_run(str(rec.get("effort_estimate", "")))
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = _GRAY_LIGHT

                if rec.get("narrative"):
                    doc.add_paragraph(rec["narrative"])

                if rec.get("recommended_actions"):
                    p = doc.add_paragraph()
                    p.add_run("Recommended Actions").bold = True
                    for action in rec["recommended_actions"]:
                        doc.add_paragraph(action, style="List Bullet")

                if rec.get("enabling_dependencies"):
                    p = doc.add_paragraph()
                    p.add_run("Must be in place first").bold = True
                    for dep in rec["enabling_dependencies"]:
                        doc.add_paragraph(dep, style="List Bullet")

                if rec.get("success_indicators"):
                    p = doc.add_paragraph()
                    p.add_run("Success Indicators").bold = True
                    for ind in rec["success_indicators"]:
                        doc.add_paragraph(ind, style="List Bullet")
    else:
        doc.add_paragraph("No recommendations available.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 6. Appendix — Full Capability List
    # ─────────────────────────────────────────────────────────────────────────
    _hdr_para(doc, "Appendix — Assessed Capabilities", level=1)

    if cap_findings:
        table = _add_table(doc, rows=len(cap_findings) + 1, cols=5,
                           col_widths=[2.2, 1.8, 1.0, 0.8, 0.8])
        _table_header_row(table, "Capability", "Domain", "Role", "Score", "Target")
        for i, c in enumerate(sorted(cap_findings, key=lambda x: x.get("domain", ""))):
            row = table.rows[i + 1]
            avg = c.get("avg_score")
            row.cells[0].text = c.get("capability_name", "")
            row.cells[1].text = c.get("domain", "")
            row.cells[2].text = c.get("capability_role", "")
            row.cells[3].text = f"{avg:.1f}" if avg is not None else "—"
            row.cells[4].text = str(c.get("target_maturity") or "—")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    else:
        doc.add_paragraph("No capability data available.")

    # ── Serialise ──────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
